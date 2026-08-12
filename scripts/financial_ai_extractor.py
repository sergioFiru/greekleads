"""
financial_ai_extractor.py — extract structured financial data from a single
document using Gemini via OpenRouter. No regex, no pdfplumber text extraction:
PDFs go straight to Gemini as a native file (it reads scanned and digital pages
alike); the exceptions are spreadsheets and Word documents, which Gemini has
no native "vision" understanding of at all — those get their text/cells
dumped to plain text first (openpyxl / python-docx), purely so there's
something for Gemini to read, not as a shortcut.

Supported: pdf, xlsx, docx natively; zip via extract_financials_from_zip()
(unzips and recurses into each supported inner file, since a zip can bundle
several fiscal years' worth of documents — same "one filing = several files"
situation merge_extractions() handles for non-zip filings). Not supported:
xls / doc (legacy pre-2007 binary Office formats) — stubbed as
not-yet-implemented since no real one has appeared in the corpus to justify
the extra dependencies (xlrd / antiword / LibreOffice conversion) that
reading them would require.

Pure extraction module — no DB, no R2, no bulk loop. Meant to be called with
raw bytes for one document at a time, e.g. by scripts/one_time/
test_financial_ai_extraction.py right now, and later by the on-demand
"Retrieve" backend for the company-page button.

Returns 13 fields — the original 6 already live in the `financial_statements`
table (fiscal_year, revenue, total_assets, equity, profit_before_tax,
net_profit) plus 7 more added 2026-08-12 (gross_profit, cost_of_sales,
inventory, cash_and_equivalents, receivables, total_liabilities,
employee_count). The 7 new ones are deliberately still standard-statutory-
statement line items (things any ΙΚΕ/ΕΠΕ/ΑΕ's normal ισολογισμός +
κατάσταση αποτελεσμάτων can show), not annual-report narrative extras like
segment/geographic revenue splits, store counts, EPS, or dividends — those
only exist for large listed groups filing full annual reports and were
deliberately excluded since they wouldn't generalize to the rest of the
corpus. All fields are nullable — a smaller company's simpler filing just
won't have some of these broken out, and null is the correct answer, not
a bug.
"""

import base64
import json
import os
import zipfile
from io import BytesIO

import requests
from docx import Document as DocxDocument
from openpyxl import load_workbook

MODEL       = "google/gemini-2.5-flash"
MAX_TOKENS  = 8192  # Gemini 2.5's "thinking" tokens eat the budget silently if too low — Scout hit this at 2048
REQUEST_TIMEOUT = 90

FIELDS = [
    "fiscal_year", "revenue", "total_assets", "equity", "profit_before_tax", "net_profit",
    "gross_profit", "cost_of_sales", "inventory", "cash_and_equivalents", "receivables",
    "total_liabilities", "employee_count",
]

PROMPT = """Είσαι ειδικός στην ανάλυση ελληνικών δημοσιευμένων οικονομικών καταστάσεων \
(ΓΕΜΗ). Το έγγραφο που ακολουθεί είναι μία δημοσιευμένη οικονομική κατάσταση \
(ισολογισμός / κατάσταση αποτελεσμάτων) μιας ελληνικής εταιρείας. Μπορεί να είναι \
αυτοματοποιημένη καταχώριση (καθαρή μορφή, πίνακες) ή σκαναρισμένο/χειρόγραφο \
έγγραφο. Μπορεί επίσης να περιέχει στήλες σύγκρισης με προηγούμενη χρήση.

Εξήγαγε ΜΟΝΟ τα στοιχεία της ΚΥΡΙΑΣ / ΤΡΕΧΟΥΣΑΣ χρήσης του εγγράφου (όχι τη στήλη \
σύγκρισης προηγούμενης χρήσης, αν υπάρχει).

Επίστρεψε ΑΠΟΚΛΕΙΣΤΙΚΑ ένα JSON object με ακριβώς αυτά τα πεδία:
{
  "found": true/false,           // false ΜΟΝΟ αν το έγγραφο δεν είναι καθόλου οικονομική κατάσταση
  "fiscal_year": 2023,           // ακέραιος, το έτος της χρήσης στην οποία αναφέρεται (όχι έτος δημοσίευσης)
  "revenue": 4650000.0,          // Κύκλος εργασιών / Πωλήσεις (καθαρός)
  "total_assets": 3300000.0,     // Σύνολο ενεργητικού
  "equity": 1520000.0,           // Σύνολο καθαρής θέσης / Ίδια κεφάλαια
  "profit_before_tax": 408000.0, // Κέρδη/ζημίες προ φόρων
  "net_profit": 318000.0,        // Καθαρά κέρδη/ζημίες (μετά φόρων)
  "gross_profit": 1250000.0,     // Μικτό αποτέλεσμα / Μικτό κέρδος (Κύκλος εργασιών - Κόστος πωλήσεων), αν αναφέρεται
  "cost_of_sales": 3400000.0,    // Κόστος πωλήσεων, ΠΑΝΤΑ ΘΕΤΙΚΟΣ αριθμός (είναι ένα κόστος, όχι μια αφαίρεση) — αν αναφέρεται ξεχωριστά
  "inventory": 610000.0,         // Αποθέματα
  "cash_and_equivalents": 220000.0, // Ταμειακά διαθέσιμα και ισοδύναμα
  "receivables": 540000.0,       // Απαιτήσεις ΣΥΝΟΛΙΚΑ: άθροισμα "Εμπορικές απαιτήσεις" + "Λοιπές απαιτήσεις" όταν εμφανίζονται ως ξεχωριστές γραμμές· αν υπάρχει ήδη ένα ενιαίο "Σύνολο απαιτήσεων", χρησιμοποίησε αυτό. Μην συμπεριλάβεις προκαταβολές για πάγια ή απαιτήσεις που ήδη περιλαμβάνονται αλλού (π.χ. στα διαθέσιμα).
  "total_liabilities": 1780000.0, // Σύνολο υποχρεώσεων (μακροπρόθεσμες + βραχυπρόθεσμες)
  "employee_count": 12,          // Αριθμός απασχολούμενου προσωπικού, αν αναφέρεται (ακέραιος)
  "notes": ""                    // σύντομη σημείωση αν κάτι είναι ασαφές ή λείπει
}

Κανόνες μορφοποίησης — ΠΟΛΥ ΣΗΜΑΝΤΙΚΟ:
- Όλες οι χρηματικές τιμές ως ΚΑΘΑΡΟΙ ΑΡΙΘΜΟΙ (JSON number, όχι string), σε ευρώ, \
χωρίς σύμβολο €, χωρίς διαχωριστικά χιλιάδων. Μετέτρεψε την ελληνική μορφή \
(π.χ. "3.300.000,00") σε 3300000.0.
- Αρνητικές τιμές (ζημίες) ως αρνητικός αριθμός (π.χ. -94000.0), όχι σε παρένθεση. \
Αυτό ισχύει ΜΟΝΟ για πεδία αποτελέσματος (net_profit, profit_before_tax, gross_profit) \
όταν είναι πράγματι ζημία — ΠΟΤΕ για cost_of_sales, το οποίο είναι πάντα θετικό ως κόστος, \
ακόμα κι αν αφαιρείται στον υπολογισμό του μικτού κέρδους.
- employee_count είναι ΑΚΕΡΑΙΟΣ (π.χ. 12), όχι χρηματική τιμή.
- Αν ένα πεδίο δεν βρίσκεται καθόλου στο έγγραφο, βάλε null — ΜΗΝ μαντέψεις.
- Αν το έγγραφο δεν είναι οικονομική κατάσταση, βάλε "found": false και όλα τα \
υπόλοιπα πεδία null.
- Απάντησε ΜΟΝΟ με το JSON object, χωρίς κανένα άλλο κείμενο.
"""


def _headers():
    key = os.environ["OPENROUTER_API_KEY"]
    return {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}


def _call_openrouter(content, model: str = MODEL) -> dict:
    """content: either a plain string, or an OpenRouter multi-part content list."""
    resp = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers=_headers(),
        json={
            "model": model,
            "messages": [{"role": "user", "content": content}],
            "max_tokens": MAX_TOKENS,
            "response_format": {"type": "json_object"},
            # Pure extraction — the right answer is fixed by the document,
            # not a creative task, so there's no reason to sample. Added
            # 2026-08-12: temperature was previously unset (provider default,
            # ~1.0 for Gemini) and run-to-run variance on the same doc/model
            # showed up repeatedly during testing.
            "temperature": 0,
        },
        timeout=REQUEST_TIMEOUT,
    )
    if resp.status_code != 200:
        return {"error": f"openrouter_http_{resp.status_code}: {resp.text[:300]}"}

    data = resp.json()
    try:
        raw = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError):
        return {"error": f"unexpected_response_shape: {json.dumps(data)[:300]}"}

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        return {"error": f"non_json_reply: {raw[:300]}"}

    result = {f: parsed.get(f) for f in FIELDS}
    result["found"] = bool(parsed.get("found", True))
    result["notes"] = parsed.get("notes") or ""
    result["error"] = None
    return result


def _xlsx_to_text(xlsx_bytes: bytes) -> str:
    wb = load_workbook(BytesIO(xlsx_bytes), data_only=True, read_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"--- Φύλλο: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _docx_to_text(docx_bytes: bytes) -> str:
    """Same rationale as _xlsx_to_text — Gemini has no native docx vision
    either, so pull out paragraph and table text and let it read that."""
    doc = DocxDocument(BytesIO(docx_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def extract_financials(file_bytes: bytes, ext: str, filename: str = "document", model: str = MODEL) -> dict:
    """
    ext: 'pdf', 'xlsx', or 'xls' (case-insensitive, leading dot optional).
    model: OpenRouter model slug — override to A/B test against a different
    model without touching the default (e.g. "google/gemini-2.5-pro").
    Returns a dict with keys: found, fiscal_year, revenue, total_assets, equity,
    profit_before_tax, net_profit, notes, error (error is None on success).
    """
    ext = ext.lower().lstrip(".")

    if ext == "pdf":
        b64 = base64.b64encode(file_bytes).decode()
        content = [
            {"type": "text", "text": PROMPT},
            {"type": "file", "file": {
                "filename": filename,
                "file_data": f"data:application/pdf;base64,{b64}",
            }},
        ]
        return _call_openrouter(content, model=model)

    if ext == "xlsx":
        text = _xlsx_to_text(file_bytes)
        return _call_openrouter(PROMPT + "\n\n--- ΠΕΡΙΕΧΟΜΕΝΟ ΥΠΟΛΟΓΙΣΤΙΚΟΥ ΦΥΛΛΟΥ ---\n" + text, model=model)

    if ext == "docx":
        text = _docx_to_text(file_bytes)
        return _call_openrouter(PROMPT + "\n\n--- ΠΕΡΙΕΧΟΜΕΝΟ ΕΓΓΡΑΦΟΥ WORD ---\n" + text, model=model)

    if ext in ("xls", "doc"):
        # Legacy binary Office formats (pre-2007) — no real doc has shown up
        # yet to test against (all 252k downloaded so far are PDF). Neither
        # openpyxl nor python-docx can read these; would need xlrd<2.0 /
        # antiword or a LibreOffice conversion step respectively. Not worth
        # building against a hypothetical — implement when a real one shows up.
        return {f: None for f in FIELDS} | {
            "found": False, "notes": "", "error": f"{ext}_not_yet_implemented",
        }

    if ext == "zip":
        # Not handled here — a zip can contain multiple documents spanning
        # multiple fiscal years, so it doesn't fit extract_financials()'s
        # "one document in, one result out" contract. Use
        # extract_financials_from_zip() instead, which returns a list of
        # already-merged per-fiscal-year results.
        return {f: None for f in FIELDS} | {
            "found": False, "notes": "", "error": "use_extract_financials_from_zip",
        }

    return {f: None for f in FIELDS} | {
        "found": False, "notes": "", "error": f"unsupported_ext:{ext}",
    }


def extract_financials_from_zip(zip_bytes: bytes, filename: str = "document.zip", model: str = MODEL) -> list[dict]:
    """
    A zip archive (e.g. a company's full-year filing bundled as one download)
    can contain several documents spanning several fiscal years — the same
    "one filing = several separate files" situation merge_extractions() was
    built for, just pre-packaged into one download instead of arriving as
    several separate financial_docs rows. Each inner file is extracted
    independently via extract_financials(), then grouped by its own
    extracted fiscal_year and merged — same "option A" strategy as
    everywhere else in this module, applied at the zip level instead of the
    company level.

    Returns a list of merge_extractions()-shaped dicts, one per distinct
    fiscal year found inside the archive (usually just one). Unsupported
    inner files (e.g. a nested .doc) are skipped, not fatal — this returns
    whatever it could get, not an all-or-nothing result.
    """
    try:
        zf = zipfile.ZipFile(BytesIO(zip_bytes))
    except zipfile.BadZipFile:
        return [{f: None for f in FIELDS} | {
            "found": False, "notes": "", "error": "bad_zip_file", "source_count": 0,
        }]

    by_fiscal_year: dict[int, list[dict]] = {}
    skipped: list[str] = []

    for name in zf.namelist():
        if name.endswith("/"):
            continue  # directory entry
        inner_ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if inner_ext not in ("pdf", "xlsx", "docx"):
            skipped.append(name)
            continue
        inner_bytes = zf.read(name)
        result = extract_financials(inner_bytes, inner_ext, filename=f"{filename}::{name}", model=model)
        if result.get("found") and not result.get("error") and result.get("fiscal_year") is not None:
            by_fiscal_year.setdefault(result["fiscal_year"], []).append(result)

    if not by_fiscal_year:
        note = f"no extractable documents found in zip (skipped: {', '.join(skipped)})" if skipped else "zip was empty or contained nothing extractable"
        return [{f: None for f in FIELDS} | {
            "found": False, "notes": note, "error": None, "source_count": 0,
        }]

    return [merge_extractions(group) for _, group in sorted(by_fiscal_year.items(), reverse=True)]


def merge_extractions(results: list[dict]) -> dict:
    """
    Merge several per-document extract_financials() results that all belong
    to the SAME fiscal year into one row.

    Why this exists: a single ΓΕΜΗ filing for one fiscal year is usually
    several separate PDFs, not one — an auditor's report, a "Προσάρτημα"
    (notes, sometimes the ONLY place employee_count shows up), a statement
    of changes in equity, the actual balance sheet/income statement, a
    management report. Each document extracts honestly on its own (nulling
    what it doesn't have, never guessing) but that means picking just one
    doc per company/year would systematically under-report. Confirmed by a
    real 15-company/45-doc test run (2026-08-12) — this is the normal case,
    not an edge case.

    Strategy (per team decision 2026-08-12, "option A"): sort candidate
    results by how complete they are (most non-null FIELDS values first),
    then fill the merged row field-by-field, first-non-null-wins — so the
    most complete single document is the primary source, and thinner
    documents only fill in gaps it left. Same strategy already agreed on
    for reconciling multiple files inside one zip archive; this generalizes
    it to "multiple docs for one fiscal year," which turns out to be the
    normal case rather than a zip-specific one.

    results: extract_financials() outputs that already share a fiscal_year
    (caller's job to group them — this function doesn't check).
    Returns one dict shaped like a single extract_financials() result, plus
    "source_count" (how many docs contributed) and "notes" (every
    contributing doc's non-empty note, concatenated, not deduplicated further
    since each usually explains a different field).
    """
    usable = [r for r in results if r.get("found") and not r.get("error")]
    if not usable:
        # Nothing usable — return the first result's shape so callers always
        # get a consistent dict back, just all-null.
        base = results[0] if results else {f: None for f in FIELDS}
        return {**{f: None for f in FIELDS}, "found": False,
                "notes": base.get("notes", ""), "error": base.get("error"), "source_count": 0}

    usable.sort(key=lambda r: sum(1 for f in FIELDS if r.get(f) is not None), reverse=True)

    merged = {f: None for f in FIELDS}
    for f in FIELDS:
        for r in usable:
            if r.get(f) is not None:
                merged[f] = r[f]
                break

    notes = [r["notes"] for r in usable if r.get("notes")]
    merged["found"] = True
    merged["notes"] = " | ".join(notes)
    merged["error"] = None
    merged["source_count"] = len(usable)
    return merged


def check_plausibility(result: dict, tolerance_pct: float = 0.03, tolerance_abs: float = 2000.0) -> list[str]:
    """
    Cheap sanity checks on a (usually merged) extraction result — built after
    a real test run (2026-08-12) caught a silent 1000x decimal-scale error
    (total_assets extracted as 60.58 instead of 60,575.64) that an LLM call
    can produce without ever signaling low confidence or erroring. These are
    accounting identities that should hold for genuinely correct figures;
    a violation doesn't prove which field is wrong, just that at least one
    of the fields involved probably is — worth a human glance, not proof of
    a bug.

    tolerance: the check tolerates the LARGER of tolerance_pct (of the
    biggest number involved) or a flat tolerance_abs (in EUR), since real
    statements often have small unreconciled roundings that aren't errors.

    Returns a list of human-readable warning strings — empty list = passed.
    """
    warnings = []

    def close_enough(a, b):
        if a is None or b is None:
            return True  # can't check what we don't have — not a failure
        diff = abs(a - b)
        return diff <= max(tolerance_abs, tolerance_pct * max(abs(a), abs(b), 1))

    assets, equity, liabilities = result.get("total_assets"), result.get("equity"), result.get("total_liabilities")
    if assets is not None and equity is not None and liabilities is not None:
        if not close_enough(assets, equity + liabilities):
            warnings.append(
                f"total_assets ({assets:,.2f}) != equity + total_liabilities "
                f"({equity:,.2f} + {liabilities:,.2f} = {equity + liabilities:,.2f})"
            )

    revenue, gross, cost = result.get("revenue"), result.get("gross_profit"), result.get("cost_of_sales")
    if revenue is not None and gross is not None and cost is not None:
        if not close_enough(revenue - cost, gross):
            warnings.append(
                f"revenue - cost_of_sales ({revenue:,.2f} - {cost:,.2f} = {revenue - cost:,.2f}) "
                f"!= gross_profit ({gross:,.2f})"
            )

    net, pbt = result.get("net_profit"), result.get("profit_before_tax")
    if net is not None and pbt is not None and net > pbt + max(tolerance_abs, tolerance_pct * max(abs(net), abs(pbt), 1)):
        # Net (after-tax) profit shouldn't exceed pre-tax profit by more than
        # rounding — if it does, a field's probably mislabeled or mis-scaled.
        warnings.append(f"net_profit ({net:,.2f}) > profit_before_tax ({pbt:,.2f}) by more than tax could explain")

    for f in ("employee_count",):
        v = result.get(f)
        if v is not None and (v < 0 or v > 200_000):  # sanity bound, not a real limit
            warnings.append(f"{f}={v} is outside a plausible range")

    return warnings
